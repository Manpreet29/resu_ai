import json
import pdfplumber
import docx
import re
import spacy
import os
import zipfile
from typing import Dict, List, Optional
from fuzzywuzzy import fuzz
from transformers import pipeline
from nltk.stem import WordNetLemmatizer
import nltk
import fitz

# Initialize NLTK
nltk.download('wordnet')
nltk.download('omw-1.4')
lemmatizer = WordNetLemmatizer()

# Load NLP model
nlp = spacy.load("en_core_web_sm")

# Initialize summarizer
try:
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
except:
    summarizer = None

# Section Keywords
WORK_KEYWORDS = [
    "work experience", "professional experience", "employment history",
    "work history", "experience", "career", "employment",
    "professional background", "positions held", "jobs"
]

EDUCATION_KEYWORDS = [
    "education", "academic background", "qualifications",
    "degrees", "university", "college", "schooling",
    "academics", "educational qualifications", "studies"
]

# Skills Database
SKILLS_DATABASE = {
    "Python", "Java", "C++", "JavaScript", "Node.js", "SQL", "React", "Docker", "Kubernetes",
    "Flask", "Django", "Git", "REST API", "GraphQL", "Pandas", "NumPy", "Scikit-learn",
    "PyTorch", "TensorFlow", "OpenCV", "Linux", "Bash", "Jenkins", "Ansible", "Azure",
    "Google Cloud", "AWS", "Spark", "Hadoop", "Tableau", "Power BI", "Agile", "Scrum", "Kanban",
    "Machine Learning", "Data Science", "Deep Learning", "Natural Language Processing", "NLP",
    "Computer Vision", "Big Data", "Artificial Intelligence", "Data Structures", "Data Analysis",
    "Data Visualization", "Statistical Modeling", "Predictive Analytics", "R"
}

def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """Extract text from PDF files"""
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"PDF Extraction Error: {e}")
        return None

def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """Extract text from DOCX files"""
    try:
        doc = docx.Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs]).strip()
    except Exception as e:
        print(f"DOCX Extraction Error: {e}")
        return None

def clean_text(text: str) -> str:
    """Clean extracted text"""
    text = re.sub(r"\(cid:\d+\)", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def extract_email(text: str) -> str:
    """Extract email addresses"""
    match = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match[0] if match else "Not Found"

def extract_phone(text: str) -> str:
    """Extract phone numbers"""
    match = re.findall(r'(\+?\d[\d\s\-().]{7,}\d)', text)
    return match[0] if match else "Not Found"

def extract_hyperlinks_from_pdf(pdf_path: str) -> Dict[str, str]:
    hyperlinks = {}
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            links = page.get_links()
            for l in links:
                uri = l.get("uri", "")
                if uri:
                    if 'linkedin.com' in uri.lower():
                        hyperlinks['linkedin'] = uri
                    elif 'github.com' in uri.lower():
                        hyperlinks['github'] = uri
    except Exception as e:
        print(f"PyMuPDF hyperlink extraction error: {e}")
    return hyperlinks

def extract_hyperlinks_from_docx(docx_path: str) -> Dict[str, str]:
    hyperlinks = {}
    try:
        # Step 1: From .rels file
        with zipfile.ZipFile(docx_path) as z:
            if 'word/_rels/document.xml.rels' in z.namelist():
                rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
                linkedin_links = re.findall(r'https?://[^\s"]*linkedin\.com[^\s"]*', rels, re.IGNORECASE)
                github_links = re.findall(r'https?://[^\s"]*github\.com[^\s"]*', rels, re.IGNORECASE)
                if linkedin_links:
                    hyperlinks['linkedin'] = linkedin_links[0]
                if github_links:
                    hyperlinks['github'] = github_links[0]
        
        # Step 2: Fallback via visible text
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        if 'linkedin' not in hyperlinks:
            linkedin_matches = re.findall(r'https?://[^\s]*linkedin\.com[^\s]*', text, re.IGNORECASE)
            if linkedin_matches:
                hyperlinks['linkedin'] = linkedin_matches[0]
        if 'github' not in hyperlinks:
            github_matches = re.findall(r'https?://[^\s]*github\.com[^\s]*', text, re.IGNORECASE)
            if github_matches:
                hyperlinks['github'] = github_matches[0]
    
    except Exception as e:
        print(f"DOCX hyperlink extraction error: {e}")
    return hyperlinks

def extract_social_links(file_path: str) -> Dict[str, str]:
    """Enhanced social link extraction with multiple fallbacks"""
    links = {}
    
    # First try hyperlink extraction
    if file_path.endswith('.pdf'):
        links.update(extract_hyperlinks_from_pdf(file_path))
    elif file_path.endswith('.docx'):
        links.update(extract_hyperlinks_from_docx(file_path))
    
    # If still not found, try advanced text pattern matching
    text = extract_text_from_pdf(file_path) if file_path.endswith('.pdf') else extract_text_from_docx(file_path)
    if text:
        # Improved LinkedIn pattern matching
        linkedin_patterns = [
            r"(?:linkedin\.com/in/[a-zA-Z0-9\-_]+)",  # standard profile URL
            r"(?:linkedin\.com/company/[a-zA-Z0-9\-_]+)",  # company pages
            r"(?:linkedin\.com/[a-zA-Z0-9\-_]+)",      # alternate profile URL
            r"linkedin\s*:\s*([a-zA-Z0-9\-_]+)",        # "LinkedIn: username" format
            r"linkedin\s*/\s*([a-zA-Z0-9\-_]+)",        # "LinkedIn/username" format
            r"li\s*:\s*([a-zA-Z0-9\-_]+)",              # "li: username" shorthand
            r"in/[a-zA-Z0-9\-_]+"                      # "/in/username" pattern
        ]
        
        # Improved GitHub pattern matching
        github_patterns = [
            r"(?:github\.com/[a-zA-Z0-9\-_]+)",        # standard profile URL
            r"github\s*:\s*([a-zA-Z0-9\-_]+)",          # "GitHub: username" format
            r"github\s*/\s*([a-zA-Z0-9\-_]+)",          # "GitHub/username" format
            r"gh\s*:\s*([a-zA-Z0-9\-_]+)",              # "gh: username" shorthand
            r"github\.io/[a-zA-Z0-9\-_]+"               # GitHub pages
        ]
        
        # Try each LinkedIn pattern
        if not links.get('linkedin'):
            for pattern in linkedin_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    username = match.group(1) if match.groups() else match.group(0)
                    if not username.startswith('linkedin.com'):
                        username = f"linkedin.com/in/{username}"
                    links['linkedin'] = f"https://{username}"
                    break
        
        # Try each GitHub pattern
        if not links.get('github'):
            for pattern in github_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    username = match.group(1) if match.groups() else match.group(0)
                    if not username.startswith('github.com'):
                        username = f"github.com/{username}"
                    links['github'] = f"https://{username}"
                    break
    
    return links

def debug_social_links(file_path: str):
    """Debug function to help identify why links aren't being found"""
    print("\n=== SOCIAL LINK DEBUGGING ===")
    
    # Show raw text extraction
    text = extract_text_from_pdf(file_path) if file_path.endswith('.pdf') else extract_text_from_docx(file_path)
    print("\nExtracted Text Sample (first 500 chars):")
    print(text[:500] + "...")
    
    # Show hyperlink extraction results
    print("\nHyperlink Extraction Results:")
    if file_path.endswith('.pdf'):
        print("PDF Hyperlinks:", extract_hyperlinks_from_pdf(file_path))
    elif file_path.endswith('.docx'):
        print("DOCX Hyperlinks:", extract_hyperlinks_from_docx(file_path))
    
    # Show text pattern matches
    print("\nText Pattern Matches:")
    linkedin_matches = re.findall(r'linkedin[^a-zA-Z0-9]*[\/:][^a-zA-Z0-9]*([a-zA-Z0-9\-_]+)', text, re.IGNORECASE)
    github_matches = re.findall(r'github[^a-zA-Z0-9]*[\/:][^a-zA-Z0-9]*([a-zA-Z0-9\-_]+)', text, re.IGNORECASE)
    print("LinkedIn patterns found:", linkedin_matches)
    print("GitHub patterns found:", github_matches)
    
def extract_skills(text: str) -> List[str]:
    """Extract skills from text"""
    text_lower = text.lower()
    extracted_skills = set()
    
    # Exact matches
    for skill in SKILLS_DATABASE:
        if re.search(rf'\b{re.escape(skill.lower())}\b', text_lower):
            extracted_skills.add(skill)
    
    # Handle abbreviations
    abbrev_mapping = {
        'nlp': 'Natural Language Processing',
        'ai': 'Artificial Intelligence',
        'ml': 'Machine Learning',
        'ds': 'Data Science'
    }
    
    for abbrev, full_skill in abbrev_mapping.items():
        if abbrev in text_lower and full_skill in SKILLS_DATABASE:
            extracted_skills.add(full_skill)
    
    return sorted(extracted_skills) if extracted_skills else ["No Skills Found"]

def extract_section(text: str, section_keywords: List[str]) -> str:
    """Extract specific sections from resume"""
    lines = text.split('\n')
    section_lines = []
    found_section = False
    
    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue
            
        # Check for section headers
        is_header = any(
            fuzz.partial_ratio(keyword.lower(), clean_line.lower()) >= 70
            for keyword in section_keywords
        ) and len(clean_line.split()) <= 5
        
        if is_header and not found_section:
            found_section = True
            continue
            
        if found_section:
            # Stop at next section
            other_sections = WORK_KEYWORDS + EDUCATION_KEYWORDS
            other_sections = [kw for kw in other_sections if kw not in section_keywords]
            is_other_header = any(
                fuzz.partial_ratio(keyword.lower(), clean_line.lower()) >= 70
                for keyword in other_sections
            )
            if is_other_header:
                break
                
            section_lines.append(clean_line)
    
    return '\n'.join(section_lines) if section_lines else "No Data"

def extract_work_experience(text: str) -> str:
    """Extract work experience section"""
    return extract_section(text, WORK_KEYWORDS)

def extract_education(text: str) -> str:
    """Extract education section"""
    return extract_section(text, EDUCATION_KEYWORDS)

def generate_resume_summary(text: str) -> str:
    """Generate AI summary of resume"""
    if not summarizer:
        return "Summary not available"
    
    # Clean text
    text = re.sub(r"\S+@\S+", "", text)  # Remove emails
    text = re.sub(r"\(?\+?\d{1,3}[-.\s]?\d{2,4}[-.\s]?\d{2,4}[-.\s]?\d{4}\)?", "", text)  # Remove phones
    
    try:
        summary = summarizer(text[:1000], max_length=100, min_length=30, do_sample=False)
        if summary and isinstance(summary, list) and 'summary_text' in summary[0]:
            return summary[0]['summary_text']
    except Exception as e:
        print(f"Summary generation error: {e}")
    
    return "Summary not available"

def parse_resume(file_path: str) -> Dict[str, str]:
    """Main resume parsing function"""
    try:
        # Debug social links first
        debug_social_links(file_path)

        # Extract social links first (from hyperlinks or text)
        social_links = extract_social_links(file_path)

        # Extract text content
        if file_path.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        else:
            return {"error": "Unsupported file format"}
            
        if not text:
            return {"error": "No readable text found"}
            
        text = clean_text(text)
        
        # Parse all sections
        parsed_data = {
            "email": extract_email(text),
            "phone": extract_phone(text),
            "linkedin": social_links.get('linkedin', 'Not Found'),
            "github": social_links.get('github', 'Not Found'),
            "skills": extract_skills(text),
            "work_experience": extract_work_experience(text),
            "education": extract_education(text),
            "summary": generate_resume_summary(text),
            "text": text
        }
        
        return parsed_data
        
    except Exception as e:
        return {"error": f"Error parsing resume: {str(e)}"}

# Example usage
if __name__ == "__main__":
    result = parse_resume("sample_resume.pdf")
    print(json.dumps(result, indent=2))